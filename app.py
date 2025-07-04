import streamlit as st
import json
import os
#from file_reader import extract_text
#from metadata_extractor import extract_metadata
import fitz
import docx
from PIL import Image
import os
import chardet
import requests


def read_txt(file_path):
 try:
    with open(file_path, 'rb') as f:
            raw_data = f.read()
            encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
    with open(file_path, 'r', encoding=encoding) as f:
        return f.read()
 except Exception as e:
        raise ValueError(f"Error reading text file: {e}")

def read_docx(file_path):
 try:
    doc1= docx.Document(file_path)
    return '\n'.join([para.text for para in doc1.paragraphs])
 except Exception as e:
        raise ValueError(f"Error reading docx: {e}")

def read_pdf(file_path):
 try:
    doc = fitz.open(file_path)
    text = ''
    for page in doc:
        text += page.get_text()
    doc.close()  
    return text
 except Exception as e:
        raise ValueError(f"Error reading PDF: {e}")
def read_image(file_path):
    try:
        with open(file_path, 'rb') as f:
            response = requests.post(
                'https://api.ocr.space/parse/image',
                files={'filename': f},
                data={
                    'apikey': 'K88180486988957',
                    'language': 'eng',
                },
            )
        result = response.json()
        if result.get("IsErroredOnProcessing"):
            raise ValueError(result.get("ErrorMessage", "OCR failed"))
        return result['ParsedResults'][0]['ParsedText']
    except Exception as e:
        raise ValueError(f"Cloud OCR failed: {e}")
def extract_text(file_path):
    ext = os.path.splitext(file_path)[-1].lower()
    if ext == '.txt':
        return read_txt(file_path)
    elif ext == '.docx':
        return read_docx(file_path)
    elif ext == '.pdf':
        return read_pdf(file_path)
    elif ext in ['.jpg', '.jpeg', '.png']:
        return read_image(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
##metadata_extractor
import spacy
from keybert import KeyBERT
from transformers import pipeline
import streamlit as st  # add this at the top
import logging
import fitz
import docx
from docx import Document

@st.cache_resource
def load_models():
    
    try:
        
        kw_model = KeyBERT('distilbert-base-nli-mean-tokens')
        summarizer = pipeline(
            "summarization", 
            model="facebook/bart-large-cnn",
            tokenizer="facebook/bart-large-cnn"
        )
        
        # Load spaCy model
        nlp = spacy.load("en_core_web_sm")
        
        return kw_model, summarizer, nlp
    except Exception as e:
        f"Error loading models: {e}"
        # Fallback to simpler models
        kw_model = KeyBERT()
        summarizer = pipeline("summarization", model="t5-small")
        nlp = spacy.load("en_core_web_sm")
        return kw_model, summarizer, nlp

    
kw_model, summarizer ,nlp= load_models()

def extract_keywords(text, top_n=10):
    keywords = kw_model.extract_keywords(text, keyphrase_ngram_range=(1, 3), stop_words='english', use_maxsum=True, nr_candidates=25)
    return [kw for kw, score in keywords]

def generate_summary(text, sentence_count=3,max_length: int = 150, min_length: int = 30):
    word_count = len(text.split())

    try:
        # If very short text, return it directly
        if word_count < 20:
            return text.strip()

        
        if word_count < 80:
            sentences = list(nlp(text).sents)
            return ' '.join(sent.text for sent in sentences[:min(sentence_count, len(sentences))])
        
        # For longer text, use BART summarizer
        summary_output = summarizer(
            text,
            max_length,
            min_length,
            do_sample=False
        )
        return summary_output[0]['summary_text']

    except Exception as e:
        return "Summary generation failed: " + str(e)
def genrate_entities(text:str):
    if not text:
        return []
    
    try:
        doc = nlp(text)
        
        # Filter out low-confidence or irrelevant entities
        entities = []
        seen_entities = set()
        
        for ent in doc.ents:
            # Skip very short entities or those with only digits
            if len(ent.text.strip()) < 2 or ent.text.strip().isdigit():
                continue
            
            # Avoid duplicates 
            entity_lower = ent.text.strip().lower()
            if entity_lower not in seen_entities:
                entities.append((ent.text.strip(), ent.label_))
                seen_entities.add(entity_lower)
        
        return entities[:20]  
    
    except Exception as e:
        (f"Entity extraction failed: {e}")
        return []

def pdf_auth(file_path):
    doc = fitz.open(file_path)
    metadata = doc.metadata
    return metadata.get("author") or metadata.get("Author")
def docx_auth(file_path):
    doc = Document(file_path)
    core_props = docx.core_properties
    return core_props.author
import re

def guess_auth(text):
    lines = text.lower().split('\n')[:20]

    # Common patterns to search
    patterns = [
        r"(?:written|created|authored|by)\s+(mr\.?\s*)?([a-z]\.?\s*){0,3}[a-z]+",  # initials + surname
        r"(?:written|created|authored|by)\s+[a-z]+(?:\s+[a-z]+){0,3}"              # full name with up to 3 words
    ]

    for line in lines:
        for pattern in patterns:
            match = re.search(pattern, line)
            if match:
                raw = match.group()
                # Clean and format result
                author = raw.replace("written", "").replace("created", "").replace("authored", "").replace("by", "")
                author = author.strip(" ,:\n\t")
                return author.title()

    # Fallback to NER
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ == "PERSON" and ent.start_char < 200:
            return ent.text

    return None
import os

# File size

def extract_title(text):
    doc = nlp(text)

    for sent in list(doc.sents)[:5]:
        cleaned = sent.text.strip()
        if 5 <= len(cleaned.split()) <= 15:
            return cleaned

    noun_chunks = [chunk.text.strip() for chunk in doc.noun_chunks if 3 <= len(chunk.text.split()) <= 10]
    if noun_chunks:
        return max(noun_chunks, key=len)

    
    return list(doc.sents)[0].text.strip() if doc.sents else "Untitled Document"

def file_info(file_path=None):
    if file_path:
        try:
            size_bytes = os.path.getsize(file_path)
            file_size = f"{size_bytes / (1024 * 1024):.2f} MB"
            file_format = os.path.splitext(file_path)[-1].lower()
        except:
            file_size = "Unknown"
            file_format = "Unknown"
    else:
        file_size = "Unknown"
        file_format = "Unknown"

    return file_size, file_format



def extract_metadata(text,file_path=None):
    doc = nlp(text)

    # Heuristic title: first sentence or longest noun chunk
    title = list(doc.sents)[0].text if doc.sents else "Untitled Document"
    #author
    author = None
    if file_path:
        if file_path.endswith(".pdf"):
            author = pdf_auth(file_path)
        elif file_path.endswith(".docx"):
            author = docx_auth(file_path)
    if not author:
        author = guess_auth(text)

    # Keywords: extract unique lowercased noun chunks (1-3 words)
    keywords = extract_keywords(text)

    # Named Entities: extract entities with labels
    entities=genrate_entities(text)
    #title
    title = extract_title(text)


   #summary
    summary_adv= generate_summary(text)
    # Return metadata dictionary
    file_size, file_format = file_info(file_path)
    metadata = {
        "title": title.strip(),
        "Author":author if author else "Unknown",
        "keywords": keywords,
        "entities": entities,
        "summary":summary_adv.strip(),
        "file_format": file_format,
        "file_size": file_size,
       
    }
    

    return metadata

# Page configuration
st.set_page_config(
    page_title="Metadata Extractor",
    page_icon="logom.png",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 10px;
        margin-bottom: 2rem;
        text-align: center;
        color: white;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    
    .success-box {
        background: #d4edda;
        color: #155724;
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #c3e6cb;
        margin: 1rem 0;
    }
    
    .warning-box {
        background: #fff3cd;
        color: #856404;
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #ffeaa7;
        margin: 1rem 0;
    }
    
    .metadata-section {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.1);
        margin: 1rem 0;
        color: black;
    }
</style>
""", unsafe_allow_html=True)

# Header
st.markdown("""
<div class="main-header">
    <h1> Metadata Extractor</h1>
    <p>Transform your documents into structured metadata using advanced AI technology</p>
</div>
""", unsafe_allow_html=True)

# Sidebar with information and settings
with st.sidebar:
   st.markdown("""
   <div style='
    background: linear-gradient(to right, #667eea, #764ba2);
    padding: 1.5rem;
    border-radius: 12px;
    color: white;
    margin-bottom: 1.5rem;
    box-shadow: 0 4px 6px rgba(0,0,0,0.1);
   '>
    <h4 style='margin-top: 0;'>📋 How it Works</h4>
    <p style='margin: 0.5rem 0;'> Step 1: Upload your document</p>
    <p style='margin: 0.5rem 0;'> Step 2: AI extracts text content</p>
    <p style='margin: 0.5rem 0;'> Step 3: Generate structured metadata</p>
    <p style='margin: 0.5rem 0;'> Step 4: Download results</p>
   </div>
   """, unsafe_allow_html=True)

    
with st.sidebar:
    st.markdown("""
    <div style='
        background: linear-gradient(to right, #667eea, #764ba2);
        padding: 1.2rem;
        border-radius: 12px;
        color: white;
        margin-bottom: 1.5rem;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    '>
        <h4 style='margin-top: 0;'>📊 Supported Formats</h4>
        <p style='margin: 0.6rem 0;'>📄 <b>PDF</b> – Portable Document Format</p>
        <p style='margin: 0.6rem 0;'>📝 <b>DOCX</b> – Microsoft Word Document</p>
        <p style='margin: 0.6rem 0;'>📋 <b>TXT</b> – Plain Text File</p>
        <p style='margin: 0.6rem 0;'>🖼️ <b>PNG/JPG</b> – Image with OCR processing</p>
    </div>
    """, unsafe_allow_html=True)



    
    st.header("⚙️ Settings")
    max_file_size = st.slider("Max file size (MB)", 1, 200, 100)
    show_text_preview = st.checkbox("Show extracted text preview", value=True)

#
st.subheader("📤 Upload Your Document")
uploaded_file = st.file_uploader(
        label= "",
        type=["pdf", "docx", "txt", "png", "jpg", "jpeg"],
        help="Select a document to extract metadata from"
    )

# Processing section
if uploaded_file is not None:
    # File info display
    file_size_mb = uploaded_file.size / (1024 * 1024)
    file_type = uploaded_file.type or "Unknown"

    file_temp = """
    <div style='padding: 10px 15px; background-color: #1e1e2f; border-radius: 8px;'>
      <p style='margin: 0; color: #aaa;'>{icon} {label}</p>
      <p style='margin: 0; font-size: 18px; font-weight: bold; color: white;'>{value}</p>
    </div>
    """
    col1, col2, col3 = st.columns(3)

    with col1:
       st.markdown(file_temp.format(icon="📁", label="File Name", value=uploaded_file.name), unsafe_allow_html=True)

    with col2:
      st.markdown(file_temp.format(icon="📏", label="File Size", value=f"{file_size_mb:.2f} MB"), unsafe_allow_html=True)

    with col3:
       st.markdown(file_temp.format(icon="📋", label="File Type", value=file_type.split('/')[-1].upper()), unsafe_allow_html=True)

    
    # File size warning
    if file_size_mb > max_file_size:
        st.markdown(f"""
        <div class="warning-box">
            <strong>⚠️ Large File Warning</strong><br>
            File size ({file_size_mb:.2f} MB) exceeds recommended limit ({max_file_size} MB). 
            Processing may take longer.
        </div>
        """, unsafe_allow_html=True)
     
    st.markdown("<div style='margin-top: 1rem;'></div>", unsafe_allow_html=True)
    
    # Process button
    if st.button("🚀 Start Processing", type="primary", use_container_width=True):
        try:
            # Save uploaded file
            file_path = os.path.join("temp_" + uploaded_file.name)
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            # Success message
            st.markdown(f"""
            <div class="success-box">
                <strong>✅ File Uploaded Successfully</strong><br>
                Ready to process: {uploaded_file.name}
            </div>
            """, unsafe_allow_html=True)
            
            # Progress tracking
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # Extract text
            status_text.text("Extracting text from document...")
            progress_bar.progress(25)
            
            text = extract_text(file_path)
            
            progress_bar.progress(50)
            status_text.text("Text extraction completed!")
            
            # Show text preview if enabled
            if show_text_preview and text:
                with st.expander("Preview Extracted Text", expanded=False):
                    preview_text = text[:1000] + "..." if len(text) > 1000 else text
                    st.text_area("Extracted Content", preview_text, height=200, disabled=True)
                    st.caption(f"Showing first 1000 characters of {len(text)} total characters")
            
            # Generate metadata
            status_text.text("Generating metadata...")
            progress_bar.progress(75)
            
            metadata = extract_metadata(text,file_path)

            
            progress_bar.progress(100)
            status_text.text("✅ Processing completed!")
            
            # Display results
            st.markdown("""
            <div class="metadata-section">
                <h3>📊 Generated Metadata</h3>
            </div>
            """, unsafe_allow_html=True)
            
            # Metadata display in tabs
            tab1, tab2, tab3 = st.tabs(["📋 Formatted View", "🔧 JSON View", "📈 Analytics"])
            
            with tab1:
              file_info_card = """
               <div style='padding: 10px 15px; background-color: #1e1e2f; border-radius: 8px; margin-bottom: 1rem;'>
                <p style='margin: 0; color: #aaa;'>📌 {label}</p>
                <p style='margin: 0; font-size: 18px; font-weight: bold; color: white;'>{value}</p>
               </div>
               """
    
              if isinstance(metadata, dict):
                for key, value in metadata.items():
                  if isinstance(value, list):
                    st.markdown(f"<h4 style='color: white;'>📌 {key.title()}</h4>", unsafe_allow_html=True)

        # Card pill style for list items
                    card_style = """
                    <div style='
                        display: flex;
                        flex-wrap: wrap;
                        gap: 10px;
                        margin-bottom: 1rem;
                     '>
                        {items}
                    </div>
                     """

                    item_template = """
                    <div style='
                     background-color: #1e1e2f;
                     color: white;
                     padding: 8px 14px;
                     border-radius: 20px;
                     font-size: 14px;
                     box-shadow: 0 2px 5px rgba(0,0,0,0.2);
                     '>{}</div>
                     """

                    items_html = ''.join([item_template.format(str(i)) for i in value])
                    st.markdown(card_style.format(items=items_html), unsafe_allow_html=True)

                  elif isinstance(value, dict):
                   st.markdown(f"<h4 style='color: white;'>📌 {key.title()}</h4>", unsafe_allow_html=True)
                   st.json(value)

                  else:
                     st.markdown(file_info_card.format(label=key.title(), value=value), unsafe_allow_html=True)

              else:
                    st.json(metadata)
            
            with tab2:
                st.json(metadata)
            
            with tab3:
                # Basic analytics
                if isinstance(metadata, dict):
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        st.markdown(file_temp.format(icon="🔢",label=" Metadata Fields", value=len(metadata)),unsafe_allow_html=True)
                    with col2:
                        st.markdown(file_temp.format(icon="📝",label= "Text Length", value=len(text)),unsafe_allow_html=True)
                    with col3:
                        word_count = len(text.split()) if text else 0
                        st.markdown(file_temp.format(icon="📖",label= "Word Count", value=word_count),unsafe_allow_html=True)
            
            # Download section
            st.markdown("📥 Download Options")
            
            col1, col2 = st.columns(2)
            
            with col1:
                # JSON download
                json_data = json.dumps(metadata, indent=2, ensure_ascii=False)
                st.download_button(
                    "Download Metadata (JSON)",
                    data=json_data,
                    file_name=f"{uploaded_file.name}_metadata.json",
                    mime="application/json",
                    use_container_width=True
                )
            
            with col2:
                # Text download
                if text:
                    st.download_button(
                        "Download Extracted Text",
                        data=text,
                        file_name=f"{uploaded_file.name}_extracted_text.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
            
            # Clean up
            try:
                os.remove(file_path)
            except:
                pass
                
        except Exception as e:
            st.error(f"❌ Error processing file: {str(e)}")
            # Clean up on error
            try:
                if 'file_path' in locals():
                    os.remove(file_path)
            except:
                pass
